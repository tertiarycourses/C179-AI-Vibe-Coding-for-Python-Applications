#!/usr/bin/env python3
"""Generate the Lesson Plan (LP) DOCX in the Tertiary house format.

Cover page + Document Version Control Record + auto TOC + Arial 11pt body +
colour-coded 3-day schedule tables (9:30am-6:30pm, 8 training hours/day, 1h
lunch, tea within, final assessment Day 3 4:00pm). Topics/labs come from
course_data + the domain data files so the LP stays aligned with the deck,
guide and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_topic1 import TOPIC1; from data_topic2 import TOPIC2
from data_topic3 import TOPIC3; from data_topic4 import TOPIC4
from data_topic5 import TOPIC5
ACT=TOPIC1+TOPIC2+TOPIC3+TOPIC4+TOPIC5
import prodoc
REPO=HERE; ASSETS=os.path.join(REPO,"courseware","assets")

BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
HEADER_FILL="1F6FEB"; TOPIC_FILL="E8F0FE"; BREAK_FILL="FFF4E5"; LUNCH_FILL="FDE9D9"; ASSESS_FILL="E8F7EE"

def lab_titles(nums):
    return "; ".join(f"Lab {a['num']}: {a['title']}" for a in ACT if a['num'] in nums)

# ------------------------------------------------ schedule (single source of truth for timing)
# (start, end, minutes, kind, activity_text)  kind: admin/topic/lab/break/lunch/assess/recap
SCHEDULE = {
 1: (C.DAY_THEMES[1], [
    ("9:30","10:00",30,"admin","Welcome, course introduction, ground rules and mandatory digital attendance (AM)"),
    ("10:00","10:45",45,"topic",f"Topic 1 — {C.TOPICS[0]['title']}: {C.TOPICS[0]['subtitle']} (concepts + demo)"),
    ("10:45","11:00",15,"break","Tea break"),
    ("11:00","13:00",120,"lab","Hands-on: "+lab_titles([1,2,3])),
    ("13:00","14:00",60,"lunch","Lunch break"),
    ("14:00","15:30",90,"lab","Hands-on: "+lab_titles([4])+f". Topic 2 — {C.TOPICS[1]['title']}: {C.TOPICS[1]['subtitle']} (concepts)"),
    ("15:30","15:45",15,"break","Tea break"),
    ("15:45","18:15",150,"lab","Hands-on: "+lab_titles([5,6,7,8])),
    ("18:15","18:30",15,"recap","Day 1 recap, Q&A and PM digital attendance"),
 ]),
 2: (C.DAY_THEMES[2], [
    ("9:30","9:45",15,"recap","Day 1 recap and mandatory digital attendance (AM)"),
    ("9:45","10:45",60,"topic",f"Topic 3 — {C.TOPICS[2]['title']}: {C.TOPICS[2]['subtitle']} (concepts + demo)"),
    ("10:45","11:00",15,"break","Tea break"),
    ("11:00","13:00",120,"lab","Hands-on: "+lab_titles([9,10])),
    ("13:00","14:00",60,"lunch","Lunch break"),
    ("14:00","15:30",90,"lab","Hands-on: "+lab_titles([11,12])),
    ("15:30","15:45",15,"break","Tea break"),
    ("15:45","18:15",150,"lab",f"Topic 4 — {C.TOPICS[3]['title']}: {C.TOPICS[3]['subtitle']} (concepts). Hands-on: "+lab_titles([13,14])),
    ("18:15","18:30",15,"recap","Day 2 recap, Q&A and PM digital attendance"),
 ]),
 3: (C.DAY_THEMES[3], [
    ("9:30","9:45",15,"recap","Day 2 recap and mandatory digital attendance (AM)"),
    ("9:45","11:00",75,"lab","Hands-on: "+lab_titles([15,16])),
    ("11:00","11:15",15,"break","Tea break"),
    ("11:15","13:00",105,"lab",f"Topic 5 — {C.TOPICS[4]['title']}: {C.TOPICS[4]['subtitle']} (concepts). Hands-on: "+lab_titles([17])),
    ("13:00","14:00",60,"lunch","Lunch break"),
    ("14:00","15:05",65,"lab","Hands-on: "+lab_titles([18,19,20])),
    ("15:05","15:20",15,"break","Tea break"),
    # Assessment block matches the v17 papers exactly: WA 50 min + CS 80 min.
    # 15:20 start keeps the day at 480 instructional minutes with the closing
    # admin block ending at 18:30.
    ("15:20","15:35",15,"assess","Briefing for Assessment"),
    ("15:35","16:25",50,"assess","Written Assessment (WA) — Short-Answer Questions (SAQ), 50 minutes, open book"),
    ("16:25","17:45",80,"assess","Case Study (CS) — a CardGuard scenario with practical tasks drawn from the labs, 80 minutes, open book"),
    ("17:45","18:30",45,"admin","TRAQOM survey, assessment digital attendance, sign Assessment Summary Record"),
 ]),
}

# ------------------------------------------------ build document
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)

prodoc.add_cover_page(doc,"LESSON PLAN",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,C.VERSIONS)
prodoc.add_toc(doc)

def H(text,level=1):
    h=doc.add_heading(text,level=level); return h

H("Course Information",1)
info=[("Course Title",C.TITLE),("WSQ Course Reference",C.COURSE_CODE),
      ("Training Provider",C.ORG+"  ("+C.UEN.replace('UEN: ','UEN ')+")"),
      ("Duration",f"{C.DAYS} days · 8 training hours per day ({C.DAYS*8} hours)"),
      ("Daily Timing","9:30 am – 6:30 pm (1-hour lunch; tea breaks within training time)"),
      ("Mode","Instructor-led, hands-on labs using uv, an AI coding assistant, VS Code and Docker"),
      ("Trainer",C.TRAINER)]
t=doc.add_table(rows=0,cols=2); t.style="Table Grid"
for k,v in info:
    c=t.add_row().cells; c[0].text=""; r=c[0].paragraphs[0].add_run(k); r.bold=True; r.font.size=Pt(10)
    prodoc._shade_cell(c[0],TOPIC_FILL)
    c[1].text=""; c[1].paragraphs[0].add_run(v).font.size=Pt(10)

H("Learning Outcomes",1)
doc.add_paragraph("On completion of this course, learners will be able to:")
for lo in C.LEARNING_OUTCOMES:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(lo).font.size=Pt(10.5)

H("Assessment",1)
for a in [C.ASSESSMENT["written"],C.ASSESSMENT["practical"],
          "Format: Open Book — course slides, Learner Guide and approved materials only.",
          "Final assessment is conducted on Day 3 from 4:00 pm.",C.ASSESSMENT["note"]]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size=Pt(10.5)

def set_cell(cell,text,bold=False,size=9.5,color=None,fill=None,align=None):
    cell.text=""; p=cell.paragraphs[0]
    if align: p.alignment=align
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.name="Arial"
    if color: r.font.color.rgb=color
    if fill: prodoc._shade_cell(cell,fill)

KIND_FILL={"topic":TOPIC_FILL,"break":BREAK_FILL,"lunch":LUNCH_FILL,"assess":ASSESS_FILL,
           "admin":"F3F5F8","recap":"F3F5F8","lab":None}

H("Course Schedule",1)
for day,(theme,rows) in SCHEDULE.items():
    H(f"Day {day} — {theme}",2)
    tbl=doc.add_table(rows=0,cols=3); tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=tbl.add_row().cells
    for i,htext in enumerate(["Time","Duration","Topic / Activity"]):
        set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
    training=0
    for start,end,mins,kind,text in rows:
        cells=tbl.add_row().cells; fill=KIND_FILL.get(kind)
        set_cell(cells[0],f"{start}–{end}",bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        set_cell(cells[1],f"{mins} min",size=9.5,fill=fill)
        set_cell(cells[2],text,bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        if kind!="lunch": training+=mins
    # widths
    for row in tbl.rows:
        row.cells[0].width=Inches(1.15); row.cells[1].width=Inches(0.9); row.cells[2].width=Inches(4.75)
    p=doc.add_paragraph(); r=p.add_run(f"Total training time: {training} minutes ({training//60} hours)."); r.italic=True; r.font.size=Pt(9.5); r.font.color.rgb=GREY
    assert training==480, f"Day {day} training minutes = {training}, expected 480"

H("Lab Reference (aligned to exam skill areas)",1)
tt=doc.add_table(rows=0,cols=3); tt.style="Table Grid"
hdr=tt.add_row().cells
for i,htext in enumerate(["Topic / Exam skill area","Weighting","Labs"]):
    set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
for tp in C.TOPICS:
    acts=[a for a in ACT if a["topic"]==tp["num"]]
    cells=tt.add_row().cells
    set_cell(cells[0],f"Topic {tp['code']}: {tp['title']}",bold=True,size=9.5,fill=TOPIC_FILL)
    set_cell(cells[1],tp["weighting"],size=9.5,fill=TOPIC_FILL)
    set_cell(cells[2],", ".join(f"Lab {a['num']}" for a in acts),size=9.5)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
OUT=os.path.join(REPO,"courseware",f"LP-{C.SHORT_TITLE}.docx")
doc.save(OUT)
print("Saved",OUT)

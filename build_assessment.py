#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""WSQ ASSESSMENT GENERATOR — WA (SAQ) + CASE STUDY (CS) variant.

Builds four DOCX for a course whose practical instrument is a single coherent
CASE STUDY (rather than discrete PP tasks):

  - Written Assessment (SAQ)  — 7 open-ended KNOWLEDGE questions (K1-K7)
  - Case Study (CS)           — one scenario + 5 questions carrying ability
                                codes (A1-A10) distributed as the original paper
  ... each as a Question Paper and a matching Answer Key.

All four carry the WSQ house cover page (via prodoc.py, same as the Lesson Plan
and Learner Guide), Arial 11 body, the copyright + page-number footer, and
EXPLICIT page breaks (never keepNext/cantSplit — see SKILL.md -> Pagination).

Runs IN PLACE from the skill directory. Override the repo with REPO=/path.
Currently configured for: Build and Deploy Python Applications with Vibe Coding.
"""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _find_repo():
    env = os.environ.get("REPO")
    if env and os.path.isdir(env):
        return os.path.abspath(env)
    d = os.path.dirname(os.path.abspath(__file__))
    while d != os.path.dirname(d):
        if os.path.isdir(os.path.join(d, ".git")) or \
           (os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "assessment"))):
            return d
        d = os.path.dirname(d)
    return os.getcwd()


REPO = _find_repo()
for _cand in (os.path.join(REPO, ".claude/skills/tertiary-lesson-plan"),
              os.path.expanduser("~/.claude/skills/tertiary-lesson-plan")):
    if os.path.exists(os.path.join(_cand, "prodoc.py")):
        sys.path.insert(0, _cand)
        break
import prodoc  # cover page + page numbers (same as LP/LG)

# ─── EDIT PER COURSE ────────────────────────────────────────────────────────
TITLE       = "Build and Deploy Python Applications with Vibe Coding"
COURSE_CODE = "TGS-2019504591"
Q_VER, A_VER = "v17", "v17"       # supersedes the retired-syllabus WA v15 / CS v16
WA_TIME     = "50 minutes"
CS_TIME     = "80 minutes"
# ────────────────────────────────────────────────────────────────────────────
prodoc.TGS = f"TGS Ref No: {COURSE_CODE}"
OUT = os.path.join(REPO, "assessment")


def _logo(name):
    here = os.path.dirname(os.path.abspath(__file__))
    for p in (os.path.join(REPO, "courseware/assets", name),
              os.path.join(REPO, "build/assets", name),
              os.path.join(here, "assets", name)):
        if os.path.exists(p):
            return p
    return None


ORG_LOGO    = _logo("tertiary-infotech-logo.png")
COURSE_LOGO = _logo("course-logo.png")   # None -> Tertiary-only cover (as LP/LG)

BRAND = RGBColor(0x1F, 0x6F, 0xEB)
DARK  = RGBColor(0x11, 0x18, 0x27)
GREY  = RGBColor(0x55, 0x5B, 0x66)

# ═══════════════════════════════════════════════ WRITTEN ASSESSMENT (KNOWLEDGE)
# Mirrors the reference SAQ exactly: 7 open-ended questions, codes K1-K7.
# Every item is answerable from the course slides (source cited in the key).
# (criterion, context, question, [model-answer points])
from assessment_data import (WRITTEN, SCENARIO_INTRO, SCENARIO_POINTS,
                             SCENARIO_TAIL, CS_ITEMS)


# ═══════════════════════════════════════════════════════════════ doc helpers
def base_doc():
    doc = Document()
    n = doc.styles["Normal"]
    n.font.name = "Arial"
    n.font.size = Pt(11)
    return doc


def para(doc, text, size=11, bold=False, italic=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align is not None:
        p.alignment = align
    return p


def heading(doc, text, size=13):
    para(doc, text, size=size, bold=True, color=BRAND, after=6, before=8)


def bullet_para(doc, text, size=11, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run("•  " + text)
    r.font.size = Pt(size)
    return p


def answer_box(doc, height_pt=110):
    """Empty bordered box giving the candidate writing space on a question paper.

    Only ever EMPTY. Model answers are NOT boxed — see model_answer(): a long
    marking guide cannot fit inside a single-row table without the row breaking
    across a page, which is precisely the defect SKILL.md -> Pagination warns
    about (Google Docs draws the border anyway and prints text through it).
    """
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    tr = t.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trh = OxmlElement('w:trHeight')
    trh.set(qn('w:val'), str(int(height_pt * 20)))
    trh.set(qn('w:hRule'), 'atLeast')
    trPr.append(trh)
    # Prevent the empty writing box from ever splitting across a page.
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)


def model_answer(doc, lines):
    """Model answer / marking guide as flowing bulleted text (never boxed)."""
    para(doc, "Suggestive answers (not exhaustive):", size=10, bold=True, after=3)
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.20)
        p.paragraph_format.space_after = Pt(2.5)
        r = p.add_run("•  " + ln)
        r.font.size = Pt(9.5)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


FILL_GAP = 6


def candidate_block(doc):
    heading(doc, "A: Trainee / Candidate Information")
    for label in ["Trainee Name (as per NRIC): ______________________________________",
                  "Last 3 digits and alphabet of NRIC/FIN: ____________________",
                  "Date: ____________________"]:
        p = para(doc, label, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0


BRIEFING = [
    "Place phones and other materials under the table or on the floor.",
    "No photos or recording of assessment scripts.",
    "No discussion during the assessment.",
    "Use a black/blue pen for hard-copy assessments.",
    "No liquid paper / correction tape.",
    "Scripts are collected when time is up.",
]

LMS_URL = "https://lms-tms.tertiaryinfotech.com/"


def add_hyperlink(p, url, text):
    """Add a real clickable Word hyperlink (blue, underlined) to paragraph p."""
    r_id = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run)
    p._p.append(link)
    return link


def instructions(doc, minutes_text, extra=None):
    heading(doc, "B: Instruction to Candidate")
    items = [
        "This is an individual exercise.",
        "This is an open-book assessment.",
        f"A total of {minutes_text} is given to complete this assessment.",
        None,  # LMS upload line
    ] + list(extra or []) + BRIEFING
    for i, s in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        if s is None:
            p.add_run(f"{i}.  Complete your answers on the document provided and "
                      "upload the completed answers to the LMS at ").font.size = Pt(11)
            add_hyperlink(p, LMS_URL, LMS_URL)
            p.add_run(".").font.size = Pt(11)
        else:
            p.add_run(f"{i}.  {s}").font.size = Pt(11)


def grading(doc, what):
    heading(doc, "Grading")
    para(doc, what, size=11, after=12)
    for ln in ["Grade: _______  (C / NYC)",
               "Assessor Name: __________________________   Assessor NRIC: ________________",
               "Date: ________________________                    Signature: ____________________"]:
        p = para(doc, ln, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0


# NOTE: the retired format put a "For Official Use Only" assessor sign-off at the
# BACK of each question paper. The house standard now carries that sign-off in the
# page-2 Grading block only — a trailing block is a duplicate and a QA failure, so
# official_use() has been removed rather than merely left uncalled.


def finish(doc, path):
    prodoc.add_page_numbers(doc)
    prodoc.enable_update_fields(doc)
    doc.save(path)
    print("  saved:", os.path.basename(path))


def doc_header(doc, kind_cover, heading_text, version):
    prodoc.add_cover_page(doc, kind_cover, TITLE, version,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, heading_text, size=13, bold=True, color=BRAND,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=12)


# ═══════════════════════════════════════════════════════════════════ builders
def build_wa(answers):
    doc = base_doc()
    kind = "Written Assessment (SAQ) — Answer Key" if answers else "Written Assessment (SAQ)"
    head = "Answers to Written Assessment (SAQ)" if answers else "Written Assessment (SAQ)"
    doc_header(doc, kind, head, A_VER if answers else Q_VER)

    if not answers:
        candidate_block(doc)
        instructions(doc, WA_TIME, extra=[
            "Answer ALL seven questions in your own words, in the box provided.",
            "All questions are based on the concepts taught in Topics 1-6.",
        ])
        grading(doc, "Candidate has answered all written questions and demonstrated the underpinning "
                     "knowledge required for the course learning outcomes.")
        page_break(doc)

    heading(doc, "C: Questions and Answers" if not answers else "C: Model Answers and Marking Guide")
    if answers:
        para(doc, "Mark each answer against the model points below. Award the mark where the candidate "
                  "covers the key concepts — wording will vary and the lists are not exhaustive. Every item "
                  "is taught in the course slides (source cited at the end of each model answer).",
             size=10.5, italic=True, color=GREY, after=8)
    else:
        para(doc, "Answer all seven questions. Each question tests underpinning knowledge covered in the "
                  "course slides. The knowledge code is shown at the end of each question.",
             size=10.5, italic=True, color=GREY, after=8)

    # EXPLICIT pagination: 2 questions per page on the paper, 1 model answer per page in the key.
    per_page = 1 if answers else 2
    for i, (crit, ctx, q, pts) in enumerate(WRITTEN, 1):
        para(doc, f"Question {i}:", size=11.5, bold=True, after=2, before=6)
        if answers:
            # Key: restate the stem compactly so each model answer fits one page.
            para(doc, ctx, size=9.5, italic=True, color=GREY, after=2)
            para(doc, q, size=10, bold=True, after=4)
            model_answer(doc, pts)
        else:
            para(doc, ctx, size=11, after=3)
            para(doc, q, size=11, bold=True, after=4)
            answer_box(doc, height_pt=145)
        if i % per_page == 0 and i < len(WRITTEN):
            page_break(doc)

    suffix = A_VER if answers else Q_VER
    name = (f"Answer to WA (SAQ) - {TITLE} - {suffix}.docx" if answers
            else f"WA (SAQ) - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))


def scenario_block(doc):
    heading(doc, "C: Case Study")
    para(doc, "Scenario — CardGuard, CardGuard Pte Ltd", size=11.5, bold=True, after=3)
    para(doc, SCENARIO_INTRO, size=11, after=6)
    for pt in SCENARIO_POINTS:
        bullet_para(doc, pt)
    para(doc, SCENARIO_TAIL, size=11, after=8, before=6)


def build_cs(answers):
    doc = base_doc()
    kind = "Case Study (CS) Assessment — Answer Key" if answers else "Case Study (CS) Assessment"
    head = "Answers to Case Study Assessment" if answers else "Case Study Assessment"
    doc_header(doc, kind, head, A_VER if answers else Q_VER)

    if not answers:
        candidate_block(doc)
        instructions(doc, CS_TIME, extra=[
            "Answer ALL five questions, in the box provided, based on the case study scenario.",
            "Where the scenario provides figures, show your working.",
            "All five questions use the techniques you practised in the hands-on labs.",
        ])
        grading(doc, "Candidate has answered all five case study questions and demonstrated the ability to "
                     "apply project management techniques to a realistic project scenario.")
        page_break(doc)

    scenario_block(doc)
    if answers:
        para(doc, "Note to assessor: each question maps to the ability criteria shown and to the labs the "
                  "candidate completed in class — the model answer is the lab procedure applied to this "
                  "scenario. Award Competent (C) where the candidate covers the substance of the model "
                  "points; exact wording, format and layout will vary.",
             size=10.5, italic=True, color=GREY, after=8)
    page_break(doc)

    # EXPLICIT pagination: each case study question gets its own page, on the paper AND in
    # the key — the prompts are long and the boxes are tall.
    for i, (label, crit, prompt, cap, lab_ref, pts) in enumerate(CS_ITEMS, 1):
        if answers:
            para(doc, f"{label} ({crit}) — Model Answer", size=12, bold=True, color=BRAND,
                 after=3, before=10)
            para(doc, prompt, size=9.5, italic=True, color=GREY, after=3)
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
            r1 = p.add_run("Reference: "); r1.bold = True; r1.font.size = Pt(10)
            r2 = p.add_run(lab_ref); r2.font.size = Pt(10)
            model_answer(doc, pts)
            # The CS marking guides are long (12-20 points each). They are UNBOXED
            # flowing text, so running on across a page boundary is safe — no border
            # can break and nothing prints through. Forcing a page break per answer
            # only produced near-empty spill pages, so the coloured
            # "Question N (Ax) — Model Answer" heading is the separator instead.
        else:
            para(doc, f"{label} ({crit}):", size=11.5, bold=True, after=3, before=4)
            para(doc, prompt, size=11, after=3)
            # The candidate sees the same lab citation as the assessor: this is an
            # open-book assessment of work they did in class, so naming the labs is
            # part of the instrument, not a hint.
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
            r1 = p.add_run("Reference: "); r1.bold = True; r1.font.size = Pt(10)
            r1.font.color.rgb = GREY
            r2 = p.add_run(lab_ref); r2.font.size = Pt(10); r2.font.color.rgb = GREY
            para(doc, cap, size=10.5, italic=True, color=GREY, after=4)
            answer_box(doc, height_pt=300)
            # Question paper: one question + its box per page, always.
            if i < len(CS_ITEMS):
                page_break(doc)

    suffix = A_VER if answers else Q_VER
    name = (f"Answer to Case Study (CS) Assessment - {TITLE} - {suffix}.docx" if answers
            else f"Case Study (CS) Assessment - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))


if __name__ == "__main__":
    os.makedirs(OUT, exist_ok=True)
    print("Building WSQ assessment set (WA + Case Study)…")
    build_wa(answers=False)
    build_wa(answers=True)
    build_cs(answers=False)
    build_cs(answers=True)
    codes = " · ".join(f"Q{i}={c}" for i, (_, c, _, _, _, _) in enumerate(CS_ITEMS, 1))
    print(f"Done. WA: {len(WRITTEN)} questions (K1-K7) · CS: {len(CS_ITEMS)} questions ({codes}).")
